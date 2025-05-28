import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document

st.title("ê³µì§€ë¬¸ ëŒ€ìƒë³„ Word ìƒì„±ê¸°")

st.warning("""
ğŸ“¢ **ê³µì§€ì‚¬í•­**

ì—…ë¡œë“œí•˜ì‹œëŠ” íŒŒì¼ì€ í•œ ì¹¸ì—ëŠ” ê¼­ **ëª¨ë‘ë³‘ì›** ë°ì´í„°ë¥¼ **í•œ ê°œì”©** ë”°ë¡œ ì°ì–´ì„œ ì…ë ¥í•´ ì£¼ì„¸ìš”.

ë°ì´í„° í˜•ì‹ì´ ë§ì§€ ì•Šìœ¼ë©´ ì²˜ë¦¬ê°€ ì–´ë ¤ìš¸ ìˆ˜ ìˆìŠµë‹ˆë‹¤.  
ì •í™•í•œ ì—…ë¡œë“œ ë¶€íƒë“œë¦½ë‹ˆë‹¤!
""")

uploaded_file = st.file_uploader("ì—‘ì…€ íŒŒì¼ì„ ì—…ë¡œë“œí•˜ì„¸ìš”", type=["xlsx"])

if uploaded_file:
    df = pd.read_excel(uploaded_file)

    required_columns = ['ë°°í¬í™€ë“œ', 'í™€ë“œì‚¬ìœ ', 'ê³µì§€ë¬¸', 'ëŒ€ìƒ']
    if not all(col in df.columns for col in required_columns):
        st.error(f"í•„ìˆ˜ ì»¬ëŸ¼ì´ ëª¨ë‘ í¬í•¨ë˜ì–´ì•¼ í•©ë‹ˆë‹¤: {required_columns}")
    else:
        def is_blank_or_dash(val):
            return pd.isna(val) or str(val).strip() in ['', '-']

        # ìœ íš¨ ë°ì´í„°ë§Œ í•„í„°ë§
        df_filtered = df[~(
            df['ë°°í¬í™€ë“œ'].apply(is_blank_or_dash) &
            df['í™€ë“œì‚¬ìœ '].apply(is_blank_or_dash) &
            df['ê³µì§€ë¬¸'].apply(is_blank_or_dash) &
            df['ëŒ€ìƒ'].apply(is_blank_or_dash)
        )]

        # 'Y' í¬í•¨ëœ í–‰ ì œê±°
        df_filtered = df_filtered[
            ~df_filtered['ë°°í¬í™€ë“œ'].astype(str).str.contains('Y', na=False) &
            ~df_filtered['í™€ë“œì‚¬ìœ '].astype(str).str.contains('Y', na=False)
        ]

        ì „ì²´ë³‘ì›_ê³µì§€ë¬¸ = []
        rows = []

        for _, row in df_filtered.iterrows():
            ëŒ€ìƒ_raw = row['ëŒ€ìƒ']
            ê³µì§€ = str(row['ê³µì§€ë¬¸']).strip()

            if pd.isna(ëŒ€ìƒ_raw):
                ëŒ€ìƒ_raw = ''

            targets = [t.strip().replace('\xa0', '') for t in str(ëŒ€ìƒ_raw).split(',') if t.strip()]

            if 'ì „ì²´ë³‘ì›' in targets:
                ì „ì²´ë³‘ì›_ê³µì§€ë¬¸.append(ê³µì§€)
                targets = [t for t in targets if t != 'ì „ì²´ë³‘ì›']

            for target in targets:
                rows.append({'ëŒ€ìƒ': target, 'ê³µì§€ë¬¸': ê³µì§€})

        # DataFrame ìƒì„± ë° ì „ì²´ë³‘ì› ê³µì§€ ë³‘í•©
        grouped_df = pd.DataFrame(rows)
        ëª¨ë“ _ëŒ€ìƒ = grouped_df['ëŒ€ìƒ'].unique().tolist()

        if grouped_df.empty and not ì „ì²´ë³‘ì›_ê³µì§€ë¬¸:
            st.warning("ìœ íš¨í•œ ë°ì´í„°ê°€ ì—†ìŠµë‹ˆë‹¤.")
        else:
            if not grouped_df.empty:
                grouped = grouped_df.groupby('ëŒ€ìƒ')['ê³µì§€ë¬¸'].apply(list).reset_index()
                # ì „ì²´ë³‘ì› ê³µì§€ë¬¸ ì¶”ê°€
                grouped['ê³µì§€ë¬¸'] = grouped['ê³µì§€ë¬¸'].apply(lambda lst: lst + ì „ì²´ë³‘ì›_ê³µì§€ë¬¸)
            else:
                grouped = pd.DataFrame({'ëŒ€ìƒ': ['ì „ì²´ë³‘ì›'], 'ê³µì§€ë¬¸': [ì „ì²´ë³‘ì›_ê³µì§€ë¬¸]})

            # ê³µì§€ë¬¸ ë¦¬ìŠ¤íŠ¸ë¥¼ ë¬¸ìì—´ë¡œ ë³€í™˜í•´ì„œ ë™ì¼í•œ ê³µì§€ë¬¸ë¼ë¦¬ ê·¸ë£¹í™”
            grouped['ê³µì§€ë¬¸_str'] = grouped['ê³µì§€ë¬¸'].apply(lambda lst: '||'.join(lst))  # êµ¬ë¶„ì ì‚¬ìš©
            merged = grouped.groupby('ê³µì§€ë¬¸_str')['ëŒ€ìƒ'].apply(list).reset_index()
            merged['ê³µì§€ë¬¸'] = merged['ê³µì§€ë¬¸_str'].apply(lambda s: s.split('||'))

            st.success("Word íŒŒì¼ ìƒì„± ì™„ë£Œ! ì•„ë˜ì—ì„œ ë‹¤ìš´ë¡œë“œí•˜ì„¸ìš”:")

            for _, row in merged.iterrows():
                ëŒ€ìƒë“¤ = sorted(row['ëŒ€ìƒ'])
                ê³µì§€ë¬¸_list = row['ê³µì§€ë¬¸']

                ëŒ€ìƒ_str = ', '.join(ëŒ€ìƒë“¤)
                doc = Document()
                doc.add_heading(f"{ëŒ€ìƒ_str} ê³µì§€ë¬¸", level=1)

                for idx, notice in enumerate(ê³µì§€ë¬¸_list, start=1):
                    doc.add_paragraph(f"{idx}. {notice}")

                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                st.download_button(
                    label=f"{ëŒ€ìƒ_str}.docx ë‹¤ìš´ë¡œë“œ",
                    data=buffer,
                    file_name=f"{ëŒ€ìƒ_str}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"download_{hash(ëŒ€ìƒ_str)}"
                )
